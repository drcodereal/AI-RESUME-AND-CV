from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate,Paragraph,Spacer
from reportlab.lib.styles import getSampleStyleSheet,ParagraphStyle
from docx import Document

def resume_pdf(r):
    bio=BytesIO(); doc=SimpleDocTemplate(bio,pagesize=A4,rightMargin=40,leftMargin=40,topMargin=35,bottomMargin=35); styles=getSampleStyleSheet(); title=ParagraphStyle("Title",parent=styles["Title"],fontSize=22,spaceAfter=6); h=ParagraphStyle("H",parent=styles["Heading2"],fontSize=12,spaceBefore=10,spaceAfter=4); body=ParagraphStyle("B",parent=styles["BodyText"],fontSize=9,leading=13)
    story=[Paragraph(r.name,title),Paragraph(f"{r.email or ''} | {r.phone or ''}",body),Paragraph(r.role or '',body)]
    for label,txt in [("SUMMARY",r.summary),("EDUCATION",r.education),("SKILLS",r.skills),("EXPERIENCE",r.experience),("PROJECTS",r.projects),("CERTIFICATIONS",r.certifications)]:
        story += [Paragraph(label,h),Paragraph((txt or "").replace("&","&amp;").replace("\n","<br/>"),body)]
    doc.build(story); bio.seek(0); return bio

def resume_docx(r):
    bio=BytesIO(); doc=Document(); doc.add_heading(r.name,0); doc.add_paragraph(f"{r.email or ''} | {r.phone or ''}"); doc.add_paragraph(r.role or '')
    for label,txt in [("SUMMARY",r.summary),("EDUCATION",r.education),("SKILLS",r.skills),("EXPERIENCE",r.experience),("PROJECTS",r.projects),("CERTIFICATIONS",r.certifications)]: doc.add_heading(label,1); doc.add_paragraph(txt or '')
    doc.save(bio); bio.seek(0); return bio
